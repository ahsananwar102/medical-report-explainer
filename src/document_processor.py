import io
import pymupdf
import pdfplumber
from docx import Document
from typing import Optional, Dict, List
import streamlit as st
from pathlib import Path

from config.config import MAX_FILE_SIZE, SUPPORTED_FORMATS

class DocumentProcessor:
    """Process and extract text from various document formats"""

    def __init__(self):
        self.supported_formats = SUPPORTED_FORMATS
        self.max_file_size = MAX_FILE_SIZE

    def validate_file(self, uploaded_file) -> Dict[str, any]:
        validation = {
            "is_valid": True,
            "errors": [],
            "file_info": {}
        }
        if uploaded_file is None:
            validation["is_valid"] = False
            validation["errors"].append("No file uploaded")
            return validation

        if uploaded_file.size > self.max_file_size:
            validation["is_valid"] = False
            validation["errors"].append(
                f"File size exceeds {self.max_file_size / (1024*1024):.1f}MB limit"
            )

        file_extension = Path(uploaded_file.name).suffix.lower()
        if file_extension not in self.supported_formats:
            validation["is_valid"] = False
            validation["errors"].append(
                f"Unsupported format. Supported: {', '.join(self.supported_formats)}"
            )

        validation["file_info"] = {
            "name": uploaded_file.name,
            "size": uploaded_file.size,
            "type": uploaded_file.type,
            "extension": file_extension
        }
        return validation

    def extract_text(self, uploaded_file) -> Dict[str, any]:
        validation = self.validate_file(uploaded_file)
        if not validation["is_valid"]:
            return {
                "success": False,
                "text": "",
                "errors": validation["errors"],
                "metadata": validation["file_info"]
            }

        ext = validation["file_info"]["extension"]
        try:
            if ext == ".pdf":
                return self._extract_pdf_text(uploaded_file)
            elif ext == ".txt":
                return self._extract_txt_text(uploaded_file)
            elif ext == ".docx":
                return self._extract_docx_text(uploaded_file)
            else:
                return {
                    "success": False,
                    "text": "",
                    "errors": [f"Extraction not implemented for {ext}"],
                    "metadata": validation["file_info"]
                }
        except Exception as e:
            return {
                "success": False,
                "text": "",
                "errors": [f"Error extracting text: {str(e)}"],
                "metadata": validation["file_info"]
            }

    def _extract_pdf_text(self, uploaded_file) -> Dict[str, any]:
        text_content = ""
        metadata = {"pages": 0, "extraction_method": "pymupdf_primary"}

        try:
            pdf_bytes = uploaded_file.read()
            uploaded_file.seek(0)
            doc = pymupdf.open(stream=pdf_bytes, filetype="pdf")
            metadata["pages"] = doc.page_count

            for page in doc:
                text_content += page.get_text() + "\n\n"
            doc.close()

            if len(text_content.strip()) < 100:
                text_content = self._extract_with_pdfplumber(uploaded_file)
                metadata["extraction_method"] = "pdfplumber_fallback"

        except Exception as e:
            try:
                text_content = self._extract_with_pdfplumber(uploaded_file)
                metadata["extraction_method"] = "pdfplumber_fallback"
            except Exception as fallback_error:
                return {
                    "success": False,
                    "text": "",
                    "errors": [f"PDF extraction failed: {e}, fallback error: {fallback_error}"],
                    "metadata": metadata
                }

        return {
            "success": True,
            "text": text_content.strip(),
            "errors": [],
            "metadata": metadata
        }

    def _extract_with_pdfplumber(self, uploaded_file) -> str:
        text_content = ""
        pdf_bytes = uploaded_file.read()
        uploaded_file.seek(0)
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for pg in pdf.pages:
                page_text = pg.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
        return text_content

    def _extract_txt_text(self, uploaded_file) -> Dict[str, any]:
        try:
            content = uploaded_file.read().decode('utf-8')
        except UnicodeDecodeError:
            uploaded_file.seek(0)
            content = uploaded_file.read().decode('latin-1', errors='ignore')
        metadata = {
            "encoding": "utf-8",
            "character_count": len(content),
            "line_count": content.count('\n')
        }
        return {"success": True, "text": content, "errors": [], "metadata": metadata}

    def _extract_docx_text(self, uploaded_file) -> Dict[str, any]:
        text_content = ""
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text_content += para.text + "\n"
        for tbl in doc.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    text_content += cell.text + " "
                text_content += "\n"
        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "character_count": len(text_content)
        }
        return {"success": True, "text": text_content.strip(), "errors": [], "metadata": metadata}

    def preprocess_medical_text(self, text: str) -> str:
        if not text:
            return ""
        lines = text.split('\n')
        cleaned = [ln.strip() for ln in lines if ln.strip()]
        cleaned_text = '\n'.join(cleaned)
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        return cleaned_text

    def extract_medical_sections(self, text: str) -> Dict[str, str]:
        sections = {key: "" for key in (
            "patient_info", "chief_complaint", "history",
            "physical_exam", "assessment", "plan",
            "medications", "other"
        )}
        patterns = {
            "patient_info": ["patient information", "demographics", "patient details"],
            ...
        }
        current = "other"
        for line in text.split('\n'):
            ll = line.lower().strip()
            for sec, pats in patterns.items():
                for pat in pats:
                    if pat in ll and len(ll) < 100:
                        current = sec
                        break
            if line.strip():
                sections[current] += line + "\n"
        for sec in sections:
            sections[sec] = sections[sec].strip()
        return sections



def create_sample_medical_text() -> str:
    """Create sample medical text for testing purposes"""
    return """
PATIENT INFORMATION:
Name: John Doe
Age: 45
Date of Birth: 01/15/1978
MRN: 123456789

CHIEF COMPLAINT:
Patient presents with chest pain and shortness of breath that started 2 hours ago.

HISTORY OF PRESENT ILLNESS:
The patient is a 45-year-old male with a history of hypertension and diabetes mellitus type 2 
who presents to the emergency department with acute onset chest pain. The pain is described as 
substernal, crushing in nature, with radiation to the left arm. Associated symptoms include 
diaphoresis, nausea, and mild dyspnea.

PHYSICAL EXAMINATION:
Vital Signs: BP 150/95, HR 102, RR 22, O2 Sat 94% on room air, Temp 98.6°F
General: Patient appears uncomfortable and diaphoretic
Cardiovascular: Regular rate and rhythm, no murmurs, rubs, or gallops
Pulmonary: Bilateral crackles at bases
Extremities: No peripheral edema

ASSESSMENT AND PLAN:
1. Acute coronary syndrome - obtain ECG, cardiac enzymes, chest X-ray
2. Hypertension - continue current medications
3. Diabetes mellitus type 2 - monitor blood glucose

MEDICATIONS:
- Lisinopril 10mg daily
- Metformin 500mg twice daily
- Aspirin 81mg daily (newly prescribed)
""" 
